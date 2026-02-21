import pdfplumber
from docx import Document
import io
import logging
from typing import Optional

logger = logging.getLogger(__name__)


class DocumentParser:
    """Parse PDF and DOCX files to extract text"""
    
    @staticmethod
    def parse_pdf(file_content: bytes) -> str:
        """Extract text from PDF file"""
        try:
            text_parts = []
            
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            
            text = "\n\n".join(text_parts)
            logger.info(f"Extracted {len(text)} characters from PDF")
            return text
            
        except Exception as e:
            logger.error(f"Error parsing PDF: {e}")
            raise ValueError(f"Failed to parse PDF: {str(e)}")
    
    @staticmethod
    def parse_docx(file_content: bytes) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(io.BytesIO(file_content))
            
            text_parts = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)
            
            text = "\n\n".join(text_parts)
            logger.info(f"Extracted {len(text)} characters from DOCX")
            return text
            
        except Exception as e:
            logger.error(f"Error parsing DOCX: {e}")
            raise ValueError(f"Failed to parse DOCX: {str(e)}")
    
    @staticmethod
    def parse_file(filename: str, file_content: bytes) -> str:
        """Parse file based on extension"""
        filename_lower = filename.lower()
        
        if filename_lower.endswith('.pdf'):
            return DocumentParser.parse_pdf(file_content)
        elif filename_lower.endswith('.docx'):
            return DocumentParser.parse_docx(file_content)
        else:
            raise ValueError(f"Unsupported file type: {filename}")
    
    @staticmethod
    def analyze_pdf_formatting(file_content: bytes) -> dict:
        """Analyze PDF for formatting issues"""
        issues = []
        risk_level = "Low"
        
        try:
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                for i, page in enumerate(pdf.pages):
                    # Check for multi-column layout
                    if page.width and page.height:
                        if page.width > page.height * 1.5:
                            issues.append(f"Page {i+1}: Possible multi-column layout detected")
                            risk_level = "High"
                    
                    # Check for tables
                    tables = page.find_tables()
                    if tables:
                        issues.append(f"Page {i+1}: Contains {len(tables)} table(s)")
                        if risk_level == "Low":
                            risk_level = "Medium"
                    
                    # Check for images
                    if hasattr(page, 'images') and page.images:
                        issues.append(f"Page {i+1}: Contains {len(page.images)} image(s)")
                        if risk_level == "Low":
                            risk_level = "Medium"
        
        except Exception as e:
            logger.warning(f"Could not analyze PDF formatting: {e}")
            issues.append("Unable to fully analyze PDF structure")
        
        return {
            "risk_level": risk_level,
            "issues": issues,
            "recommendations": DocumentParser._get_recommendations(risk_level)
        }
    
    @staticmethod
    def _get_recommendations(risk_level: str) -> list:
        """Get recommendations based on risk level"""
        if risk_level == "High":
            return [
                "Convert to single-column layout",
                "Remove tables and use plain text lists",
                "Remove images and icons",
                "Use simple, ATS-friendly formatting"
            ]
        elif risk_level == "Medium":
            return [
                "Consider simplifying complex layouts",
                "Ensure text is extractable from tables",
                "Test with different ATS systems"
            ]
        else:
            return [
                "Resume appears ATS-friendly",
                "Continue using current format"
            ]
